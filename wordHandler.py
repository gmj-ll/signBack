from docx import Document
from docx.shared import Inches
import re
import sys
import json


doc = Document(sys.argv[1])

# parag_num = 0
# for para in doc.paragraphs :
#     print(para.text)
#     parag_num  += 1  
# print ('This document has ', parag_num , ' paragraphs')

paragraphs = []

# 按段落读取全部数据
for paragraph in doc.paragraphs:
  text = paragraph.text.replace('\n', '').replace('\r', '').replace(' ', '').replace('：', '').replace(':', '')
  if text != '':
    paragraphs.append(text)

celltext = []

# 按表格读取全部数据
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
          text = cell.text.replace('\n', '').replace('\r', '').replace(' ', '').replace('：', '').replace(':', '')
          if text != '':
            celltext.append(text)

content = {}

content['title'] = sorted(set(paragraphs), key=paragraphs.index)
content['cell'] = sorted(set(celltext), key=celltext.index)
print (json.dumps(content))
