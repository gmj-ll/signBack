# from docx import Document
# from docx.shared import Inches
# # 按表格读取全部数据

# doc = Document('demo.docx')
# for table in doc.tables:
#     for row in table.rows:
#       content = row.cells[0].text.replace('\n', '').replace('\r', '').replace(' ', '').replace('：', '').replace(':', '') 
#       # print('key:' + content)
#       # print('value1:' + row.cells[1].text)
#       i = 0
#       for cell in row.cells:
#         i = i + 1
#         print('第'+ str(i) + '行:' + cell.text)
#         print('-----------')

from docx import Document
from docx.shared import Inches
import re
import sys
import json
import re
import time

fileName = 'demo.docx'
doc = Document(fileName)
obj = {'姓名':'','性别':'男','出生年月':'','政治面貌':'','文化程度':'','行政职务':'','专业技术职务':'','工作岗位':'','个人总结':'','部门考核意见':'','审批意见': '','个人确认':''}

hasSet = set()
localtime = time.localtime(time.time())
year = localtime.tm_year
mon = localtime.tm_mon
day = localtime.tm_mday

for p in doc.paragraphs:
  for run in p.runs:
    print(run.text)
# 按表格读取全部数据
# for table in doc.tables:
#     for row in table.rows:
      
#       matchKey = ''
#       for cell in row.cells:
        
        
#         i = -1
#         for paragraph in cell.paragraphs:
#           i = i + 1
#           print('第' + str(i) + '行：')
#           print(paragraph.text)
#         if re.match('.*年.*月.*日', cell.text):
#           # cell.text = re.sub('\d*年', str(year) + '年', cell.text)
#           # cell.text = re.sub('\d*月', str(mon) + '月', cell.text)
#           # cell.text = re.sub('\d*日', str(day) + '日', cell.text)
#           # print('---------')
#           print(cell[0])
#         content = cell.text.replace('\n', '').replace('\r', '').replace(' ', '').replace('：', '').replace(':', '')
#         if content in obj:
#           i = i + 1 
#           matchKey = content
#           continue
#         if not (matchKey in hasSet):
#           cell.text = cell.text.replace(cell.text, obj[matchKey])
#           hasSet.add(matchKey)
        
#         if re.match('.*年.*月.*日', cell.text):
#           cell.text = re.sub('\d*年', str(year) + '年', cell.text)
#           cell.text = re.sub('\d*月', str(mon) + '月', cell.text)
#           cell.text = re.sub('\d*日', str(day) + '日', cell.text)
#           print(cell.text)
