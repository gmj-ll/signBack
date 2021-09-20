from docx import Document
from docx.shared import Inches
import re
import sys
import json
import time
from docx import shared
fileName = sys.argv[1]
doc = Document(sys.argv[1])
obj = json.loads(sys.argv[2])
title = json.loads(sys.argv[3])

hasSet = set()

localtime = time.localtime(time.time())
year = localtime.tm_year
mon = localtime.tm_mon
day = localtime.tm_mday

for p in doc.paragraphs:
  content = p.text.replace('\n', '').replace('\r', '').replace(' ', '').replace('：', '').replace(':', '')
  if content in title:
    p.text = p.text + title[content]

# 按表格读取全部数据
for table in doc.tables:
    for row in table.rows:
      i = -1
      matchKey = ''
      for cell in row.cells:
        content = cell.text.replace('\n', '').replace('\r', '').replace(' ', '').replace('：', '').replace(':', '')
        if content in obj:
          i = i + 1 
          matchKey = content
          continue
        if not (matchKey in hasSet):
          index = 0
          for paragraph in cell.paragraphs:
            index = index + 1
            if index == 1:
              if re.match('./public/', obj[matchKey]):
                addRun = paragraph.add_run()
                addRun.add_picture(obj[matchKey], width = shared.Inches(1))
              else:
                paragraph.text = paragraph.text.replace(paragraph.text, obj[matchKey])
            elif re.match('.*年.*月.*日', paragraph.text):
              paragraph.text = re.sub('\d*年', str(year) + '年', paragraph.text)
              paragraph.text = re.sub('\d*月', str(mon) + '月', paragraph.text)
              paragraph.text = re.sub('\d*日', str(day) + '日', paragraph.text)
            else:
              paragraph.clear()
          hasSet.add(matchKey)

doc.save('./public/' + fileName.replace('./uploads/', ''))

content = {}

content['msg'] = '修改成功'
# content['cell'] = sorted(set(celltext), key=celltext.index)

print(json.dumps(content))